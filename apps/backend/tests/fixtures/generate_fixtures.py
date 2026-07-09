"""Sinh CV mẫu (fixture) cho test Parser — chạy: `uv run python tests/fixtures/generate_fixtures.py`.

Tạo dữ liệu giả-thực-tế (NFR-4: KHÔNG dùng CV thật/cá nhân). Bạn có thể tự thả CV thật vào thư
mục này để thử thêm với /api/agents/parse-cv. PDF dùng ký tự ASCII (font base-14 không có dấu
tiếng Việt); DOCX giữ tiếng Việt đầy đủ (đọc thẳng XML, không phụ thuộc font).
"""

from __future__ import annotations

from pathlib import Path

import fitz  # PyMuPDF
from docx import Document

FIXTURES = Path(__file__).resolve().parent

_GOOD_PDF = """NGUYEN VAN A
Senior Backend Engineer
Email: nguyenvana@example.com | Phone: +84 901 234 567

PROFESSIONAL SUMMARY
Backend engineer with 5 years building scalable async APIs in Python.

SKILLS
Python, FastAPI, PostgreSQL, Docker, LangGraph, AWS, Redis

EXPERIENCE
Acme Corp - Senior Backend Engineer (2021 - 2024)
  Built async microservices handling 2M requests/day.
Beta Ltd - Backend Engineer (2019 - 2021)
  Designed REST APIs and CI/CD pipelines.

EDUCATION
Hanoi University of Science and Technology - B.Sc. Computer Science (2015 - 2019)
"""

_SPARSE_PDF = """TRAN VAN B
Email: tranvanb@example.com
Looking for a junior developer position. Available immediately.
"""

_NOT_A_CV = """The quick brown fox jumps over the lazy dog. This document describes the migratory
patterns of arctic terns across the northern hemisphere and has nothing to do with
employment, hiring, or resumes. Lorem ipsum dolor sit amet, consectetur adipiscing elit,
sed do eiusmod tempor incididunt ut labore et dolore magna aliqua.
"""

# CV có CHỨNG CHỈ + NGÔN NGỮ + GIẢI THƯỞNG + khối lạ (Hobbies -> other). Dùng cho slice 01c:
# kiểm parser xếp chứng chỉ vào certificates (KHÔNG vào other), hobby vào other.
_CERT_PDF = """NGUYEN VAN C
Backend Developer
Email: nguyenvanc@example.com | Phone: +84 900 111 222

SKILLS
Node.js, Express.js, MongoDB, REST API

EXPERIENCE
DevHouse - Backend Intern (2023 - 2024)
  Built REST APIs with Node.js and Express.

EDUCATION
FPT University - B.Sc. Software Engineering (2019 - 2023)

CERTIFICATES
TOEIC 945/990 (2025)
AWS Certified Cloud Practitioner (2024)

LANGUAGES
English - Professional working proficiency
Japanese - Basic (N5)

AWARDS
First prize, University Hackathon 2022

HOBBIES
Chess, reading tech blogs, hiking
"""


def _write_pdf(name: str, text: str) -> None:
    doc = fitz.open()
    page = doc.new_page()
    page.insert_textbox(fitz.Rect(50, 50, 545, 792), text, fontsize=11, fontname="helv")
    doc.save(FIXTURES / name)
    doc.close()


def _write_good_docx() -> None:
    doc = Document()
    doc.add_heading("Nguyễn Văn A", level=0)
    doc.add_paragraph("Kỹ sư Backend cao cấp | Senior Backend Engineer")
    doc.add_paragraph("Email: nguyenvana@example.com | Điện thoại: +84 901 234 567")
    doc.add_heading("Tóm tắt", level=1)
    doc.add_paragraph("5 năm kinh nghiệm xây dựng API bất đồng bộ với Python.")
    doc.add_heading("Kỹ năng", level=1)
    doc.add_paragraph("Python, FastAPI, PostgreSQL, Docker, LangGraph, AWS, Redis")
    doc.add_heading("Kinh nghiệm", level=1)
    doc.add_paragraph("Acme Corp — Senior Backend Engineer (2021–2024): vi dịch vụ async 2M req/ngày.")
    doc.add_paragraph("Beta Ltd — Backend Engineer (2019–2021): thiết kế REST API và CI/CD.")
    doc.add_heading("Học vấn", level=1)
    doc.add_paragraph("ĐH Bách khoa Hà Nội — Cử nhân Khoa học Máy tính (2015–2019).")
    doc.save(FIXTURES / "good_cv.docx")


def main() -> None:
    _write_good_docx()
    _write_pdf("good_cv.pdf", _GOOD_PDF)
    _write_pdf("sparse_cv.pdf", _SPARSE_PDF)
    _write_pdf("not_a_cv.pdf", _NOT_A_CV)
    _write_pdf("cert_cv.pdf", _CERT_PDF)
    names = [p.name for p in sorted(FIXTURES.glob("*cv*.*")) if p.suffix in {".pdf", ".docx"}]
    print("Generated fixtures:", ", ".join(names))


if __name__ == "__main__":
    main()
