"""Test Parser (slice-01) — LLM MOCK (không gọi API thật, giữ make test nhanh + không tốn credit).

Phủ: parse CV mẫu (PDF+DOCX) với LLM mock -> lưu đúng + confidence cao; file rỗng/đuôi lạ -> parse_failed;
LLM ném exception -> parse_failed (không sập); trích text THẬT trên fixture (kiểm extractor, không mock).
Ref: plan slice-01 §3.7, PRD §7.1.
"""

from __future__ import annotations

import os
from pathlib import Path

import pytest
from docx import Document

from app.agents.nodes import parser as parser_mod
from app.agents.nodes.parser import _confidence, parse_cv, parser_node
from app.core.config import settings
from app.services.storage import StorageNotFound
from app.schemas.parsed_cv import (
    Certificate,
    Education,
    Experience,
    Language,
    OtherItem,
    ParsedCV,
)
from app.tools.cv_reader import EmptyCVTextError, extract_text

FIXTURES = Path(__file__).resolve().parent / "fixtures"


def _full_parsed() -> ParsedCV:
    return ParsedCV(
        full_name="Nguyễn Văn A",
        email="nguyenvana@example.com",
        phone="+84901234567",
        skills=["Python", "FastAPI", "PostgreSQL"],
        experiences=[Experience(company="Acme Corp", title="Senior Backend Engineer")],
        education=[Education(school="ĐH Bách khoa Hà Nội", degree="Cử nhân")],
        total_years_experience=5.0,
        professional_summary="Backend engineer.",
    )


class _FakeLLM:
    """Giả ChatOpenAI.with_structured_output: .invoke(prompt) -> ParsedCV cố định."""

    def __init__(self, parsed: ParsedCV) -> None:
        self._parsed = parsed

    def invoke(self, _prompt: str) -> ParsedCV:
        return self._parsed


class _BoomLLM:
    def invoke(self, _prompt: str):
        raise RuntimeError("OpenAI API down")


# ── parse_cv với LLM mock trên CV mẫu thật (extract_text chạy thật) ──────────────
# Slice 06: parse_cv nhận BYTES + tên (chọn bộ đọc theo đuôi) — không còn mở path.


def _fixture(name: str) -> bytes:
    return (FIXTURES / name).read_bytes()


def test_parse_good_docx_mocked() -> None:
    result = parse_cv(_fixture("good_cv.docx"), "good_cv.docx", llm=_FakeLLM(_full_parsed()))
    assert result["uncertainty_flags"] == []
    assert result["confidence"] == 1.0  # đủ 5 trường lõi
    assert result["escalation_reason"] is None
    assert result["parsed_data"]["full_name"] == "Nguyễn Văn A"
    assert "Python" in result["parsed_data"]["skills"]


def test_parse_good_pdf_mocked() -> None:
    result = parse_cv(_fixture("good_cv.pdf"), "good_cv.pdf", llm=_FakeLLM(_full_parsed()))
    assert result["uncertainty_flags"] == []
    assert result["confidence"] == 1.0
    assert result["parsed_data"]["email"] == "nguyenvana@example.com"


def test_parse_works_with_storage_key_name() -> None:
    """Tên truyền vào có thể là KEY storage (cv/1/<uuid>.pdf) — vẫn chọn đúng bộ đọc theo đuôi."""
    result = parse_cv(_fixture("good_cv.pdf"), "cv/1/abc123.pdf", llm=_FakeLLM(_full_parsed()))
    assert result["uncertainty_flags"] == []
    assert result["parsed_data"]["full_name"] == "Nguyễn Văn A"


def test_confidence_reflects_completeness() -> None:
    # Chỉ có tên + email -> 2/5 trường lõi = 0.4.
    sparse = ParsedCV(full_name="Tran Van B", email="b@example.com")
    result = parse_cv(_fixture("sparse_cv.pdf"), "sparse_cv.pdf", llm=_FakeLLM(sparse))
    assert result["confidence"] == 0.4
    assert result["uncertainty_flags"] == []


# ── parse_failed: đuôi lạ / text rỗng / lỗi LLM ─────────────────────────────────


def test_unsupported_extension_parse_failed() -> None:
    # không truyền llm: extract_text raise TRƯỚC khi build LLM
    result = parse_cv(b"noi dung khong phai pdf docx " * 5, "resume.txt")
    assert result["uncertainty_flags"] == ["parse_failed"]
    assert result["confidence"] == 0.0
    assert result["escalation_reason"]
    assert result["parsed_data"] is None


def test_empty_text_parse_failed(tmp_path: Path) -> None:
    tiny = tmp_path / "tiny.docx"
    doc = Document()
    doc.add_paragraph("Hi")  # < 50 ký tự -> EmptyCVTextError
    doc.save(str(tiny))
    result = parse_cv(tiny.read_bytes(), "tiny.docx")
    assert result["uncertainty_flags"] == ["parse_failed"]
    assert result["confidence"] == 0.0


def test_llm_error_parse_failed_no_crash() -> None:
    result = parse_cv(_fixture("good_cv.pdf"), "good_cv.pdf", llm=_BoomLLM())
    assert result["uncertainty_flags"] == ["parse_failed"]
    assert result["confidence"] == 0.0
    assert "LLM" in result["escalation_reason"]


# ── extract_text THẬT (không mock) — kiểm extractor PDF/DOCX từ BYTES ────────────


def test_extract_text_real_docx() -> None:
    text = extract_text(_fixture("good_cv.docx"), "good_cv.docx")
    assert len(text.strip()) >= 50
    assert "Python" in text


def test_extract_text_real_pdf() -> None:
    text = extract_text(_fixture("good_cv.pdf"), "good_cv.pdf")
    assert len(text.strip()) >= 50
    assert "Backend" in text


def test_extract_text_short_raises(tmp_path: Path) -> None:
    tiny = tmp_path / "tiny.docx"
    doc = Document()
    doc.add_paragraph("xyz")
    doc.save(str(tiny))
    try:
        extract_text(tiny.read_bytes(), "tiny.docx")
        raise AssertionError("phải ném EmptyCVTextError")
    except EmptyCVTextError:
        pass


# ── parser_node: stub khi ENABLE_LLM=false; thật khi bật + có key CV ────────────
# Slice 06: node là ASYNC và lấy bytes qua seam storage (mock ở đây — không chạm đĩa/R2).


class _FakeStorage:
    """Storage giả: trả bytes cố định cho mọi key (hoặc ném lỗi để test nhánh hỏng)."""

    def __init__(self, data: bytes | None = None, error: Exception | None = None) -> None:
        self._data = data
        self._error = error
        self.asked: list[str] = []

    async def get(self, key: str) -> bytes:
        self.asked.append(key)
        if self._error:
            raise self._error
        return self._data or b""


async def test_parser_node_stub_when_llm_disabled(monkeypatch) -> None:
    # Ép enable_llm=False (độc lập .env) -> giữ stub (không phá flow cũ).
    monkeypatch.setattr(settings, "enable_llm", False)
    out = await parser_node({"input": {"cv_path": "cv/1/x.docx"}, "scratchpad": {}})
    assert out["confidence"] == 1.0
    assert "[parser] stub" in out["messages"][0]
    assert "parsed_data" not in out


async def test_parser_node_real_reads_via_storage(monkeypatch) -> None:
    monkeypatch.setattr(settings, "enable_llm", True)
    monkeypatch.setattr(parser_mod, "_build_parser_llm", lambda: _FakeLLM(_full_parsed()))
    storage = _FakeStorage(_fixture("good_cv.docx"))
    monkeypatch.setattr(parser_mod, "get_storage", lambda: storage)

    out = await parser_node({"input": {"cv_path": "cv/7/abc.docx"}, "scratchpad": {}})

    assert storage.asked == ["cv/7/abc.docx"]  # ĐỌC QUA SEAM, không mở path
    assert out["confidence"] == 1.0
    assert out["uncertainty_flags"] == []
    assert out["parsed_data"]["full_name"] == "Nguyễn Văn A"
    assert "[parser] OK" in out["messages"][0]


async def test_parser_node_storage_error_parse_failed_no_crash(monkeypatch) -> None:
    """Mất file / R2 lỗi / key cũ → parse_failed + escalation, KHÔNG sập pipeline (PRD §7.1)."""
    monkeypatch.setattr(settings, "enable_llm", True)
    storage = _FakeStorage(error=StorageNotFound("không tìm thấy"))
    monkeypatch.setattr(parser_mod, "get_storage", lambda: storage)

    out = await parser_node({"input": {"cv_path": "cv/9/missing.pdf"}, "scratchpad": {}})

    assert out["uncertainty_flags"] == ["parse_failed"]
    assert out["confidence"] == 0.0
    assert "storage" in out["escalation_reason"].lower()


# ── slice 01c: certificates / languages / awards / other ────────────────────────


def _parsed_with_extras() -> ParsedCV:
    return ParsedCV(
        full_name="Nguyen Van C",
        email="c@example.com",
        skills=["Node.js"],
        certificates=[Certificate(name="TOEIC", detail="945/990", year="2025")],
        languages=[Language(name="English", proficiency="Professional working")],
        awards=["First prize, University Hackathon 2022"],
        other=[OtherItem(label="Hobbies", content="Chess, reading tech blogs, hiking")],
    )


def test_new_blocks_passthrough() -> None:
    pd = parse_cv(_fixture("cert_cv.pdf"), "cert_cv.pdf", llm=_FakeLLM(_parsed_with_extras()))["parsed_data"]
    assert pd["certificates"][0] == {"name": "TOEIC", "detail": "945/990", "year": "2025"}
    assert pd["languages"][0]["name"] == "English"
    assert pd["awards"] == ["First prize, University Hackathon 2022"]
    assert pd["other"][0]["label"] == "Hobbies"


def test_certificate_not_in_other_shape() -> None:
    # Đúng SHAPE ưu tiên: chứng chỉ ở certificates, khối lạ ở other, chứng chỉ KHÔNG lẫn other.
    pd = parse_cv(_fixture("cert_cv.pdf"), "cert_cv.pdf", llm=_FakeLLM(_parsed_with_extras()))["parsed_data"]
    assert any(c["name"] == "TOEIC" for c in pd["certificates"])
    other_blob = " ".join(f"{o['label']} {o['content']}" for o in pd["other"]).lower()
    assert "toeic" not in other_blob
    assert "chess" in other_blob  # khối 'Hobbies' nằm ở other


def test_backward_compat_empty_new_blocks() -> None:
    # CV cũ (mock không set trường mới) -> [] mặc định; confidence GIỮ NGUYÊN (5 khối lõi).
    result = parse_cv(_fixture("good_cv.docx"), "good_cv.docx", llm=_FakeLLM(_full_parsed()))
    pd = result["parsed_data"]
    assert pd["certificates"] == [] and pd["languages"] == [] and pd["awards"] == [] and pd["other"] == []
    assert result["confidence"] == 1.0


def test_confidence_ignores_new_blocks() -> None:
    # name + email = 2/5; certificates/languages/awards/other KHÔNG cộng vào confidence.
    p = ParsedCV(
        full_name="X", email="x@y.com",
        certificates=[Certificate(name="TOEIC", detail="990")],
        languages=[Language(name="English")], awards=["a"],
        other=[OtherItem(label="l", content="c")],
    )
    assert _confidence(p) == 0.4


@pytest.mark.skipif(not os.environ.get("RUN_PARSE_IT"), reason="cần RUN_PARSE_IT=1 + OPENAI_API_KEY")
def test_certificates_extracted_real() -> None:
    # LLM THẬT trên cert_cv.pdf: TOEIC vào certificates, KHÔNG lọt other (kiểm ưu tiên thật).
    pd = parse_cv(_fixture("cert_cv.pdf"), "cert_cv.pdf")["parsed_data"]
    assert "toeic" in " ".join(c["name"] for c in pd["certificates"]).lower()
    other_blob = " ".join(f"{o['label']} {o['content']}" for o in pd["other"]).lower()
    assert "toeic" not in other_blob
